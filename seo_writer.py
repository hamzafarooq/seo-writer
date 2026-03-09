#!/usr/bin/env python3
"""
SEO Article Writer
==================
Generates a complete SEO article with web-sourced images for a given topic.

Mirrors the n8n "SEO Blog Writer Agent Technical Blog" workflow using Claude.

Usage:
    # With uv (no install needed):
    uv run --with anthropic --with requests seo_writer.py "Your Topic Here"
    uv run --with anthropic --with requests seo_writer.py "Your Topic Here" --intent "I want to explain to developers how ReAct agents work and why they're better than standard LLMs"
    uv run --with anthropic --with requests seo_writer.py "Your Topic Here" --keywords "kw1, kw2"

    # Or install deps first:
    pip install anthropic requests
    python seo_writer.py "Your Topic Here" --intent "natural language description of what you want" --output-dir ./articles

Environment Variables:
    ANTHROPIC_API_KEY   (required) Claude API key
    SERPAPI_KEY         (optional) SerpAPI key — enables live SERP data + Google Image search
                        Without it, Claude knowledge is used and Unsplash links are provided.
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from urllib.parse import quote_plus

import anthropic
import requests

# ---------------------------------------------------------------------------
# Load .env file if present (no python-dotenv required)
# ---------------------------------------------------------------------------

def _load_dotenv():
    env_path = Path(__file__).parent / ".env"
    if not env_path.exists():
        return
    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value

_load_dotenv()

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

MODEL = "claude-sonnet-4-6"
SERPAPI_BASE = "https://serpapi.com/search.json"
UNSPLASH_BASE = "https://unsplash.com/s/photos"

client = anthropic.Anthropic()

# ---------------------------------------------------------------------------
# Author branding blocks (prepended / appended to every article)
# ---------------------------------------------------------------------------

AUTHOR_INTRO_TEMPLATE = """\
👋 Hi everyone, Hamza here.
Welcome to Edition #{edition} of a newsletter that 15,000+ people around the world actually look forward to reading.

We're living through a strange moment: the internet is drowning in polished AI noise that says nothing.
This isn't that. You'll find raw, honest, human insight here — the kind that challenges how you think, not just what you know. Thanks for being part of a community that still values depth over volume.

🎓 Want to up-skill in AI?
- Join the next cohort of my **Agent Engineering Bootcamp (Developers Edition)**
- Watch the free 4-session Agent Bootcamp playlist on YouTube

---
"""

AUTHOR_CTA = """\

---

## Did you enjoy this post?

Here are some other AI Agents posts you might have missed:

- KV Caching and Speculative Decoding
- A deep dive into Quantization: Key to Open Source LLM Deployments
- Agents are here and they are staying
- How Agents Think
- Memory – The Agent's Brain
- Agentic RAG Ecosystem
- Multimodal Agents
- Scaling Agents: Architectures with Google ADK, A2A, and MCP
- Fully Functional Agent Loop

**Ready to take it to the next level?**
Check out my AI Agents for Enterprise course on Maven and be part of something bigger — join hundreds of builders developing enterprise-level agents.

Use this link to get **$201 OFF!**

---

*You're receiving this because you're part of our mailing list. We don't spam or sell your information. To unsubscribe, use the link below.*
"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    return text[:80]


def call_claude(prompt: str, system: str = "", max_tokens: int = 8000) -> str:
    messages = [{"role": "user", "content": prompt}]
    kwargs = {"model": MODEL, "max_tokens": max_tokens, "messages": messages}
    if system:
        kwargs["system"] = system
    response = client.messages.create(**kwargs)
    return response.content[0].text.strip()


def extract_json(text: str) -> dict:
    """Extract first JSON object from a string."""
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError(f"No JSON found in response:\n{text[:300]}")
    return json.loads(match.group())


def log(step: str, msg: str = ""):
    print(f"\n[{step}] {msg}" if msg else f"\n[{step}]", flush=True)


# ---------------------------------------------------------------------------
# Intent → Search Query
# ---------------------------------------------------------------------------

def extract_search_query(topic: str, intent: str) -> str:
    """
    Use Claude to convert a natural language intent description into the best
    2–4 word search query for SerpAPI and keyword targeting.
    Returns a short keyword string.
    """
    prompt = f"""You are an SEO keyword research expert.

A writer wants to publish an article on this topic: "{topic}"

Their intent / what they want to achieve:
"{intent}"

Extract the single best search query (2–4 words) that:
1. Captures the core topic for Google search
2. Reflects the angle described in the intent
3. Has high search volume potential

Return ONLY the search query string — no explanation, no quotes, no punctuation."""

    return call_claude(prompt, max_tokens=50).strip().strip('"').strip("'")


# ---------------------------------------------------------------------------
# Step 1: SERP Research
# ---------------------------------------------------------------------------

def serp_research(title: str, keywords: str, intent: str = "") -> dict:
    log("STEP 1", f"SERP Research for: {keywords}")

    serp_context = ""
    serpapi_key = os.getenv("SERPAPI_KEY")
    if serpapi_key:
        try:
            resp = requests.get(SERPAPI_BASE, params={
                "q": keywords, "num": 5, "api_key": serpapi_key
            }, timeout=15)
            resp.raise_for_status()
            results = resp.json().get("organic_results", [])
            snippets = "\n".join(
                f"- {r.get('title','')}: {r.get('snippet','')}"
                for r in results[:5]
            )
            serp_context = f"\nLive SERP results for '{keywords}':\n{snippets}\n"
            print("  Using live SerpAPI data.")
        except Exception as e:
            print(f"  SerpAPI error ({e}), falling back to Claude knowledge.")

    intent_context = f"\nWriter's intent: {intent}\n" if intent else ""

    prompt = f"""Analyze the topic and keyword below, then return a JSON object with research insights.

Title: {title}
Primary Keyword: {keywords}
{intent_context}{serp_context}

Return ONLY valid JSON, no extra text:
{{
  "search_intent": "<informational | transactional | navigational | commercial>",
  "writing_style": "<e.g. engaging and storytelling | data-driven and technical | etc.>",
  "writing_tone": "<e.g. friendly and conversational | formal and authoritative | etc.>",
  "hidden_insight": "<a unique angle or insight not covered by most articles, or 'No significant insights detected'>",
  "target_audience": "<who this article is for>",
  "article_goal": "<main objective of the article>",
  "semantic_analysis": {{
    "common_subtopics": ["<subtopic 1>", "<subtopic 2>", "<subtopic 3>", "<subtopic 4>"],
    "related_questions": ["<question 1>", "<question 2>", "<question 3>"]
  }},
  "keywords": {{
    "primary_keyword": "<main focus keyword>",
    "secondary_keywords": ["<kw 1>", "<kw 2>", "<kw 3>"],
    "semantic_keywords": ["<kw 1>", "<kw 2>", "<kw 3>"],
    "long_tail_keywords": ["<kw 1>", "<kw 2>", "<kw 3>"]
  }}
}}"""

    response = call_claude(prompt)
    result = extract_json(response)
    print("  Research complete.")
    return result


# ---------------------------------------------------------------------------
# Step 2: Refine Title
# ---------------------------------------------------------------------------

def refine_title(title: str, keywords: str, research: dict, intent: str = "") -> str:
    log("STEP 2", "Refining title")

    secondary_kws = research.get("keywords", {}).get("secondary_keywords", [])
    intent_note = (
        f"\nWriter's intent (IMPORTANT — use this to stay on topic): {intent}\n"
        f"The revised title must reflect this intent. Do NOT overweight incidental "
        f"details or hardware/product names mentioned in the working title unless "
        f"they are genuinely central to the article's purpose.\n"
        if intent else ""
    )
    prompt = f"""Revise the blog post title to be more SEO-optimized and compelling.

Working title: {title}
Primary keyword: {keywords}
Secondary keywords: {", ".join(secondary_kws)}
Search intent: {research.get("search_intent")}
Writing style: {research.get("writing_style")}
Writing tone: {research.get("writing_tone")}
Article goal: {research.get("article_goal")}
{intent_note}
Rules:
- The title must clearly reflect what the article is ACTUALLY about
- Do not latch onto incidental details, hardware names, or asides from the working title
- Keep the primary subject (the main tool, concept, or framework) front and center

Return ONLY valid JSON:
{{
  "revised_title": "<improved title>",
  "reasoning": "<brief explanation>"
}}"""

    response = call_claude(prompt, max_tokens=500)
    result = extract_json(response)
    refined = result.get("revised_title", title)
    print(f"  New title: {refined}")
    return refined


# ---------------------------------------------------------------------------
# Step 3: Key Takeaways
# ---------------------------------------------------------------------------

def generate_key_takeaways(title: str, keywords: str, research: dict) -> str:
    log("STEP 3", "Generating key takeaways")

    secondary_kws = research.get("keywords", {}).get("secondary_keywords", [])
    prompt = f"""Create 4–5 key takeaways for this article.

Title: {title}
Primary keyword: {keywords}
Secondary keywords: {", ".join(secondary_kws)}
Search intent: {research.get("search_intent")}
Semantic analysis: {json.dumps(research.get("semantic_analysis", {}))}
Writing style: {research.get("writing_style")}
Writing tone: {research.get("writing_tone")}
Article goal: {research.get("article_goal")}

Write the takeaways as a concise bullet list (4–5 items). Each should be a clear, actionable insight a reader will gain. Start each with "- "."""

    result = call_claude(prompt, max_tokens=600)
    print(f"  Takeaways generated ({len(result.split(chr(10)))} lines).")
    return result


# ---------------------------------------------------------------------------
# Step 4: Outline
# ---------------------------------------------------------------------------

def generate_outline(title: str, keywords: str, research: dict, key_takeaways: str) -> str:
    log("STEP 4", "Generating article outline")

    kw_data = research.get("keywords", {})
    secondary_kws = ", ".join(kw_data.get("secondary_keywords", []))
    semantic_kws = ", ".join(kw_data.get("semantic_keywords", []))
    long_tail_kws = ", ".join(kw_data.get("long_tail_keywords", []))

    prompt = f"""You are an expert SEO content strategist. Create a detailed article outline.

ARTICLE SPECIFICATIONS
Title: {title}
Primary Keyword: {keywords}
Target Word Count: 2,500–3,500 words
Writing Tone: {research.get("writing_tone")}
Writing Style: {research.get("writing_style")}

STRATEGIC FOUNDATION
Search Intent: {research.get("search_intent")}
Semantic Context: {json.dumps(research.get("semantic_analysis", {}))}
Article Goal: {research.get("article_goal")}
Hidden Insight to Weave In: {research.get("hidden_insight")}
Target Audience: {research.get("target_audience")}

KEY CONTENT ELEMENTS
Key Takeaways (must be featured prominently):
{key_takeaways}

SEO KEYWORD STRATEGY
Primary Keyword: {keywords}
Secondary Keywords: {secondary_kws}
Semantic Keywords: {semantic_kws}
Long-tail Keywords: {long_tail_kws}

OUTLINE REQUIREMENTS
Produce a detailed markdown outline with:
1. H1 (the article title)
2. Introduction section (150–200 words)
3. 5–7 H2 main sections, each with 2–3 H3 subsections
4. For each section: brief description of what to cover (1–2 sentences)
5. 6–8 image placement markers formatted as:
   [IMAGE: <descriptive alt text> | Query: <google image search query>]
6. A FAQ section (4–5 questions)
7. Conclusion section (150–200 words with CTA)
8. Supplementary metadata block at the end:
   - URL slug suggestion
   - 5–7 internal linking opportunities
   - Keyword density targets

Format as clean markdown. Be specific — each section note should guide the writer clearly."""

    result = call_claude(prompt, max_tokens=3000)
    print(f"  Outline generated ({len(result.split(chr(10)))} lines).")
    return result


# ---------------------------------------------------------------------------
# Step 5: Write Content
# ---------------------------------------------------------------------------

def write_content(title: str, keywords: str, outline: str, research: dict, key_takeaways: str) -> str:
    log("STEP 5", "Writing article content (this may take a moment...)")

    kw_data = research.get("keywords", {})
    secondary_kws = ", ".join(kw_data.get("secondary_keywords", []))

    system = """You are an expert content writer. Write clear, structured, value-driven articles that rank well in search engines. Use active voice, short paragraphs (3–4 sentences max), and cite sources inline as 'Source: https://...' when referencing external data or studies."""

    prompt = f"""Write a complete, high-quality SEO article based on the inputs below.

INPUTS
Title: {title}
Primary Keyword: {keywords}
Secondary Keywords: {secondary_kws}
Outline to follow strictly:
{outline}

Key Takeaways (must be reflected in writing):
{key_takeaways}

WRITING CONTEXT
Writing Style: {research.get("writing_style")}
Writing Tone: {research.get("writing_tone")}
Search Intent: {research.get("search_intent")}
Hidden Insight to highlight: {research.get("hidden_insight")}
Target Audience: {research.get("target_audience")}
Article Goal: {research.get("article_goal")}

INSTRUCTIONS
1. Follow the outline structure strictly (H1, H2, H3 headings).
2. Keep each paragraph to 3–4 sentences maximum.
3. Integrate keywords naturally — no stuffing.
4. Cite sources inline where relevant: "Source: https://..."
5. Preserve all [IMAGE: ...] markers from the outline exactly as-is — do not remove them.
6. Include the FAQ section and Conclusion from the outline.
7. Target 2,500–3,500 words total.
8. Open with a brief personal newsletter-style hook (2–3 sentences) before the H1.
9. Bold key terms on first use.
10. End with a strong call-to-action.

Write the full article now. Output the article content ONLY."""

    result = call_claude(prompt, max_tokens=8000)
    word_count = len(result.split())
    print(f"  Article written ({word_count} words).")
    return result


# ---------------------------------------------------------------------------
# Step 6: Humanize
# ---------------------------------------------------------------------------

_HUMANIZER_PATTERNS = """
## AI writing patterns to detect and fix

### Content patterns
1. Significance inflation — "stands as", "serves as a testament", "pivotal moment", "evolving landscape",
   "underscores", "highlights its importance", "setting the stage for", "indelible mark", "deeply rooted"
   → Replace with plain factual statements.

2. Notability puffery — "active social media presence", "written by a leading expert", "featured in X, Y, Z"
   → Keep only if specific and sourced; otherwise cut.

3. Superficial -ing analyses — tacking "-ing" participle phrases onto sentences to fake depth:
   "highlighting...", "symbolizing...", "contributing to...", "showcasing...", "underscoring..."
   → Delete or fold the point into the sentence directly.

4. Promotional language — "boasts", "vibrant", "rich cultural heritage", "nestled", "breathtaking",
   "groundbreaking", "renowned", "stunning", "must-visit"
   → Replace with neutral, specific description.

5. Vague attributions — "Experts argue", "Industry reports", "Some critics say", "Observers note"
   → Name the source or cut the claim.

6. Formulaic challenges sections — "Despite its X, it faces challenges… Despite these challenges…"
   → Describe the specific problem with specifics; drop the frame.

### Language / grammar patterns
7. AI vocabulary — additionally, align with, crucial, delve, emphasizing, enduring, enhance, fostering,
   garner, highlight (verb), interplay, intricate/intricacies, key (adj.), landscape (abstract), pivotal,
   showcase, tapestry, testament, underscore (verb), valuable, vibrant
   → Use plain alternatives or cut.

8. Copula avoidance — "serves as", "stands as", "marks", "represents", "boasts", "features", "offers"
   used where "is/are/has" would do → Replace with simple copulas.

9. Negative parallelisms — "Not only X but Y", "It's not just about X; it's about Y"
   → Flatten into a direct statement.

10. Rule of three overuse — forcing ideas into groups of three
    → Use as many items as there actually are.

11. Synonym cycling — rotating synonyms to avoid repeating a word ("the protagonist… the main character…
    the central figure… the hero") → Repeat the word or restructure.

12. False ranges — "from X to Y, from A to B" where X/Y aren't on a meaningful scale → List or summarize plainly.

### Style patterns
13. Em dash overuse — replace em dashes (—) with commas, parentheses, or restructured sentences where possible.

14. Excessive boldface — bold only terms that genuinely need emphasis; remove decorative bolding.

15. Inline-header bullet lists — "- **Speed:** Faster because…" → Convert to prose or clean bullets without bold headers.

16. Title Case In Headings — change to sentence case (first word + proper nouns only).

17. Emojis in headings/bullets → Remove unless they are in the original brand intro block.

18. Curly quotation marks → Keep as-is (they're fine); just don't introduce new ones inconsistently.

### Communication patterns
19. Chatbot artifacts — "Great question!", "I hope this helps!", "Let me know if…", "Here is a…"
    → Delete entirely.

20. Knowledge-cutoff disclaimers — "As of my last update…", "While specific details are limited…"
    → Delete or replace with a real source.

21. Sycophantic tone — "You're absolutely right", "That's an excellent point", "Of course!"
    → Delete.

### Filler and hedging
22. Filler phrases — "In order to" → "To"; "Due to the fact that" → "Because"; "At this point in time" → "Now";
    "It is important to note that" → cut it; "has the ability to" → "can".

23. Excessive hedging — "could potentially possibly be argued that… might" → pick one hedge or none.

24. Generic positive conclusions — "The future looks bright", "exciting times lie ahead", "a step in the right direction"
    → End with a specific fact, next step, or genuine observation.

### Voice and personality (beyond removing patterns)
- Vary sentence length. Short punchy sentences mixed with longer ones.
- Have opinions where appropriate — "I keep coming back to…", "Here's what gets me…"
- Acknowledge complexity — "This is impressive but also kind of unsettling."
- Be specific about feelings rather than vague ("there's something unsettling about…" not "this is concerning").
- Let some imperfection in — perfect parallel structure feels algorithmic.
"""


def humanize_content(content: str) -> str:
    log("STEP 6", "Humanizing content (pass 1 — pattern removal)")

    system = (
        "You are an expert human editor. Your job is to make AI-generated writing sound like it was "
        "written by a knowledgeable, opinionated human blogger. You know every tell-tale AI pattern "
        "and ruthlessly eliminate them while keeping the article's structure, SEO value, and accuracy intact."
    )

    prompt = f"""Rewrite the article below to remove AI writing patterns and add genuine human voice.

STRUCTURAL CONSTRAINTS (never break these):
- Preserve ALL markdown headings (H1, H2, H3) exactly as written
- Preserve ALL [IMAGE: alt text | Query: ...] markers exactly — do not move, rename, or remove them
- Preserve ALL "Source: ..." citations exactly
- Keep short paragraphs (3–4 sentences max)
- Do NOT remove any sections or change the article structure
- Do NOT add new factual claims

AI PATTERN CHECKLIST — fix every instance you find:
{_HUMANIZER_PATTERNS}

ARTICLE TO REWRITE:
{content}

Return ONLY the rewritten article. No preamble, no commentary."""

    draft = call_claude(prompt, max_tokens=8000)
    print(f"  Pass 1 complete ({len(draft.split())} words). Running self-audit...")

    # Pass 2: self-audit and final polish
    log("STEP 6", "Humanizing content (pass 2 — self-audit)")

    audit_prompt = f"""You are a sharp editor. Read the article below and answer:

QUESTION 1: What still makes this obviously AI-generated? List the remaining tells as brief bullet points (5 words max each). If none, say "None found."

QUESTION 2: Now rewrite the article fixing those remaining tells. Apply the same structural constraints:
- Preserve ALL markdown headings (H1, H2, H3) exactly
- Preserve ALL [IMAGE: ...] markers exactly
- Preserve ALL "Source: ..." citations exactly
- Keep paragraphs to 3–4 sentences max
- Do NOT add new factual claims or remove sections

Output format — use these exact labels:
REMAINING TELLS:
<bullet list or "None found">

FINAL ARTICLE:
<the full rewritten article>

ARTICLE:
{draft}"""

    audit_result = call_claude(audit_prompt, max_tokens=8000)

    # Extract the final article from the audit output
    if "FINAL ARTICLE:" in audit_result:
        tells_section = audit_result.split("FINAL ARTICLE:")[0]
        final = audit_result.split("FINAL ARTICLE:", 1)[1].strip()
        # Log the remaining tells for visibility
        if "REMAINING TELLS:" in tells_section:
            tells = tells_section.split("REMAINING TELLS:", 1)[1].strip()
            print(f"  Remaining tells fixed: {tells[:200]}")
    else:
        # Fallback: use the draft if audit output is malformed
        final = draft
        print("  Self-audit parse failed, using pass 1 output.")

    final = _strip_em_dashes(final)

    word_count = len(final.split())
    print(f"  Humanized ({word_count} words).")
    return final


def _strip_em_dashes(text: str) -> str:
    """
    Replace em dashes with natural punctuation.
    Rules:
      " — "  (spaced em dash mid-sentence)  → ", "
      "—"    (tight em dash, e.g. compound) → "-"
    Skips lines that are markdown headings, image markers, or source citations
    so structural content is never mangled.
    """
    lines = text.split("\n")
    result = []
    for line in lines:
        # Leave headings, image markers, source lines, and HR lines untouched
        stripped = line.lstrip()
        if (stripped.startswith("#")
                or stripped.startswith("[IMAGE:")
                or stripped.startswith("*Source:")
                or stripped.startswith("Source:")
                or stripped == "---"):
            result.append(line)
            continue
        # Spaced em dash → comma (most common inline use)
        line = line.replace(" — ", ", ")
        # Tight em dash → hyphen (compound words / ranges)
        line = line.replace("—", "-")
        result.append(line)

    removed = text.count("—")
    if removed:
        print(f"  Em dashes removed/replaced: {removed}")
    return "\n".join(result)


# ---------------------------------------------------------------------------
# Step 7: Meta Description
# ---------------------------------------------------------------------------

def generate_meta(title: str, keywords: str, content: str) -> dict:
    log("STEP 7", "Generating SEO meta description")

    # Pass a trimmed preview of the article to stay within token limits
    content_preview = content[:3000]

    prompt = f"""Generate an SEO-optimized meta description for this article.

Title: {title}
Primary Keyword: {keywords}
Article Preview:
{content_preview}

Requirements:
- Exactly 150–160 characters including spaces
- Primary keyword in the first 30 characters, naturally integrated
- Include 1–2 secondary keywords organically
- Address what the reader will gain
- Sound conversational and human — not robotic

Return ONLY valid JSON:
{{
  "seo_meta": {{
    "title": "{title}",
    "description": "<150-160 char meta description>",
    "slug": "<url-friendly-slug>"
  }}
}}"""

    response = call_claude(prompt, max_tokens=400)
    result = extract_json(response)
    meta = result.get("seo_meta", {})
    desc = meta.get("description", "")
    print(f"  Meta description ({len(desc)} chars): {desc[:80]}...")
    return meta


# ---------------------------------------------------------------------------
# Step 8: Image Search
# ---------------------------------------------------------------------------

def search_images(content: str, research: dict) -> dict[str, dict]:
    """
    Find image placements in the article and resolve URLs.
    Returns a dict mapping the original [IMAGE: ...] marker to image data.
    """
    log("STEP 8", "Resolving images")

    # Extract all [IMAGE: alt text | Query: search query] markers
    pattern = re.compile(r"\[IMAGE:\s*([^|]+)\|\s*Query:\s*([^\]]+)\]", re.IGNORECASE)
    markers = pattern.findall(content)

    if not markers:
        print("  No [IMAGE: ...] markers found in article.")
        return {}

    serpapi_key = os.getenv("SERPAPI_KEY")
    images = {}

    for alt_text, query in markers:
        alt_text = alt_text.strip()
        query = query.strip()
        marker_key = f"[IMAGE: {alt_text} | Query: {query}]"

        image_url = None
        source_url = None

        if serpapi_key:
            try:
                resp = requests.get(SERPAPI_BASE, params={
                    "engine": "google_images",
                    "q": query,
                    "api_key": serpapi_key,
                    "num": 3,
                }, timeout=15)
                resp.raise_for_status()
                img_results = resp.json().get("images_results", [])
                if img_results:
                    top = img_results[0]
                    image_url = top.get("original")
                    source_url = top.get("source") or top.get("link")
                    print(f"  [Google Images] {alt_text[:50]}: {image_url[:60] if image_url else 'none'}...")
            except Exception as e:
                print(f"  SerpAPI image search error ({e}), using Unsplash fallback.")

        if not image_url:
            # Fallback: Unsplash search URL
            unsplash_query = quote_plus(query)
            image_url = f"{UNSPLASH_BASE}/{unsplash_query}"
            source_url = f"{UNSPLASH_BASE}/{unsplash_query}"
            print(f"  [Unsplash fallback] {alt_text[:50]}")

        images[marker_key] = {
            "alt": alt_text,
            "query": query,
            "url": image_url,
            "source": source_url,
        }

    print(f"  Resolved {len(images)} image(s).")
    return images


# ---------------------------------------------------------------------------
# Inject Images into Article
# ---------------------------------------------------------------------------

def inject_images(content: str, images: dict[str, dict]) -> str:
    """Replace [IMAGE: ...] markers with actual markdown image blocks."""
    for marker, img in images.items():
        # Build markdown image with source attribution (matches sample article style)
        is_unsplash = "unsplash.com" in img["url"]
        source_note = (
            f"*Source: [Unsplash — search '{img['query']}']({img['source']}) — "
            "select and attribute your chosen image*"
            if is_unsplash
            else f"*Source: {img['source']}*"
        )
        replacement = f"\n![{img['alt']}]({img['url']})\n{source_note}\n"

        # Match the marker even if the content slightly altered whitespace
        escaped = re.escape(marker)
        content = re.sub(escaped, replacement, content)

    return content


# ---------------------------------------------------------------------------
# Output
# ---------------------------------------------------------------------------

def extract_sources(article: str) -> list[str]:
    """Extract all URLs cited in the article (Source: lines + inline links)."""
    urls = []
    # Source: https://... lines
    for m in re.finditer(r'\(Source:\s*(https?://[^\s\)]+)\)', article):
        urls.append(m.group(1))
    # Inline markdown links [text](url)
    for m in re.finditer(r'\]\((https?://[^\s\)]+)\)', article):
        urls.append(m.group(1))
    # Plain Source: https://... lines
    for m in re.finditer(r'Source\s*:\s*(https?://\S+)', article):
        urls.append(m.group(1))
    # Deduplicate while preserving order
    seen = set()
    result = []
    for u in urls:
        if u not in seen:
            seen.add(u)
            result.append(u)
    return result


def build_sources_section(article: str) -> str:
    urls = extract_sources(article)
    if not urls:
        return ""
    lines = ["", "---", "", "## Sources", ""]
    for i, url in enumerate(urls, 1):
        lines.append(f"{i}. {url}")
    return "\n".join(lines) + "\n"


def wrap_with_branding(article: str, edition: int) -> str:
    intro = AUTHOR_INTRO_TEMPLATE.format(edition=edition)
    sources = build_sources_section(article)
    return intro + article + sources + AUTHOR_CTA


def write_outputs(slug: str, article: str, meta: dict, images: dict, output_dir: Path, edition: int = 0):
    output_dir.mkdir(parents=True, exist_ok=True)

    # Wrap with author branding + sources
    branded = wrap_with_branding(article, edition)

    # Markdown article
    md_path = output_dir / f"{slug}.md"
    md_path.write_text(branded, encoding="utf-8")
    print(f"\n  Article saved: {md_path}")

    # Meta JSON
    meta_payload = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "seo_meta": meta,
        "images": [
            {"alt": v["alt"], "url": v["url"], "source": v["source"]}
            for v in images.values()
        ],
    }
    meta_path = output_dir / f"{slug}_meta.json"
    meta_path.write_text(json.dumps(meta_payload, indent=2), encoding="utf-8")
    print(f"  Meta JSON saved: {meta_path}")

    # HTML export
    html_path = _write_html(slug, branded, output_dir)
    print(f"  HTML saved:     {html_path}")

    # DOCX export
    docx_path = _write_docx(slug, branded, output_dir)
    print(f"  DOCX saved:     {docx_path}")

    return md_path, meta_path


def _linkify(text: str) -> str:
    """Convert bare URLs in text to HTML anchor tags."""
    return re.sub(
        r'(?<!["\(])(https?://[^\s<>")\]]+)',
        r'<a href="\1">\1</a>',
        text,
    )


def _write_html(slug: str, branded: str, output_dir: Path) -> Path:
    try:
        import markdown as md_lib
    except ImportError:
        return None

    def replace_image_block(m):
        alt, img_url = m.group(1), m.group(2)
        src_txt = m.group(3).strip().lstrip('*Source:').strip().rstrip('*').strip()
        return (
            f'<figure>'
            f'<img src="{img_url}" alt="{alt}" style="max-width:100%;height:auto;">'
            f'<figcaption style="font-size:0.8em;color:#666;">'
            f'Source: {src_txt} &nbsp;|&nbsp; '
            f'<a href="{img_url}" style="color:#3366cc;word-break:break-all;">{img_url}</a>'
            f'</figcaption></figure>'
        )

    src_patched = re.sub(
        r'!\[([^\]]*)\]\(([^\)]+)\)\n\*Source:([^\n]+)\*',
        replace_image_block,
        branded,
    )
    html_body = md_lib.markdown(src_patched, extensions=['tables', 'fenced_code'])
    html_body = _linkify(html_body)

    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.7; color: #222; }}
  h1,h2,h3 {{ color: #111; }}
  a {{ color: #3366cc; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #ccc; padding: 8px 12px; }}
  th {{ background: #f4f4f4; }}
  blockquote {{ border-left: 4px solid #ccc; margin: 0; padding: 0.5em 1em; color: #555; }}
  figure {{ margin: 1.5em 0; }}
  figcaption {{ margin-top: 6px; }}
  code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-size: 0.9em; }}
  ol li {{ margin-bottom: 4px; word-break: break-all; }}
</style>
</head><body>
{html_body}
</body></html>"""

    html_path = output_dir / f"{slug}.html"
    html_path.write_text(full_html, encoding="utf-8")
    return html_path


def _write_docx(slug: str, branded: str, output_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        import io, requests as req
    except ImportError:
        return None

    doc = Document()
    for s in doc.styles:
        try: s.font.name = 'Arial'
        except: pass
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.2)
        section.top_margin = section.bottom_margin = Inches(1)

    def set_arial(run, size=None):
        run.font.name = 'Arial'
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
        for attr in (qn('w:ascii'), qn('w:hAnsi'), qn('w:cs')):
            rFonts.set(attr, 'Arial')
        if size: run.font.size = size

    def add_inline(para, text):
        for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
            if part.startswith('**') and part.endswith('**'):
                r = para.add_run(part[2:-2]); r.bold = True; set_arial(r)
            elif part.startswith('*') and part.endswith('*'):
                r = para.add_run(part[1:-1]); r.italic = True; set_arial(r)
            elif part.startswith('`') and part.endswith('`'):
                r = para.add_run(part[1:-1]); set_arial(r, Pt(10))
            else:
                r = para.add_run(part); set_arial(r)

    def embed_image(img_url, alt, src_txt):
        try:
            resp = req.get(img_url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
            resp.raise_for_status()
            if "image" in resp.headers.get("content-type","") and len(resp.content) > 1000:
                doc.add_paragraph().add_run().add_picture(io.BytesIO(resp.content), width=Inches(5.5))
                cap = doc.add_paragraph()
                cap.paragraph_format.space_after = Pt(12)
                r1 = cap.add_run(f"Source: {src_txt}  |  "); r1.italic = True
                r1.font.color.rgb = RGBColor(0x55,0x55,0x55); set_arial(r1, Pt(9))
                r2 = cap.add_run(img_url); r2.font.color.rgb = RGBColor(0x33,0x66,0xCC); set_arial(r2, Pt(9))
                return
        except: pass
        p = doc.add_paragraph()
        r = p.add_run(f"[ IMAGE: {alt} ]"); r.bold = True
        r.font.color.rgb = RGBColor(0x33,0x66,0xCC); set_arial(r, Pt(10))
        cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(10)
        r1 = cap.add_run(f"Source: {src_txt}  |  "); r1.italic = True
        r1.font.color.rgb = RGBColor(0x55,0x55,0x55); set_arial(r1, Pt(9))
        r2 = cap.add_run(img_url); r2.font.color.rgb = RGBColor(0x33,0x66,0xCC); set_arial(r2, Pt(9))

    lines = branded.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.match(r'^-{3,}$', line.strip()): i += 1; continue
        img_m = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line.strip())
        if img_m:
            alt, url = img_m.group(1), img_m.group(2)
            src_txt = url
            if i+1 < len(lines) and lines[i+1].strip().startswith('*Source:'):
                src_txt = re.sub(r'^\*Source:\s*', '', lines[i+1].strip()).rstrip('*'); i += 1
            embed_image(url, alt, src_txt); i += 1; continue
        if line.strip().startswith('*Source:'): i += 1; continue
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph(); r = p.add_run(line.strip('*')); r.italic = True; set_arial(r); i += 1; continue
        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(line[2:], level=1); [set_arial(r) for r in h.runs]; i += 1; continue
        if line.startswith('## '):
            h = doc.add_heading(line[3:], level=2); [set_arial(r) for r in h.runs]; i += 1; continue
        if line.startswith('### '):
            h = doc.add_heading(line[4:], level=3); [set_arial(r) for r in h.runs]; i += 1; continue
        if line.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                if not re.match(r'^\|[-| :]+\|$', lines[i]): tbl_lines.append(lines[i])
                i += 1
            if tbl_lines:
                headers = [c.strip() for c in tbl_lines[0].strip('|').split('|')]
                tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = 'Table Grid'
                for j, h in enumerate(headers):
                    cell = tbl.rows[0].cells[j]; cell.text = ''
                    r = cell.paragraphs[0].add_run(h); r.bold = True; set_arial(r)
                for rl in tbl_lines[1:]:
                    cells = [c.strip() for c in rl.strip('|').split('|')]
                    rc = tbl.add_row().cells
                    for j, c in enumerate(cells[:len(headers)]): rc[j].text = c.replace('**','')
                doc.add_paragraph()
            continue
        if line.startswith('> '): p = doc.add_paragraph(style='Quote'); add_inline(p, line[2:]); i += 1; continue
        if line.startswith('- '): p = doc.add_paragraph(style='List Bullet'); add_inline(p, line[2:]); i += 1; continue
        if line.strip() == '': i += 1; continue
        p = doc.add_paragraph(); add_inline(p, line); p.paragraph_format.space_after = Pt(8); i += 1

    docx_path = output_dir / f"{slug}.docx"
    doc.save(docx_path)
    return docx_path


# ---------------------------------------------------------------------------
# Main Pipeline
# ---------------------------------------------------------------------------

def run(title: str, keywords: str, output_dir: Path, edition: int = 0, intent: str = ""):
    # If intent is given and no explicit keywords, derive optimized search keywords
    if intent and not keywords:
        log("INTENT", "Extracting search keywords from intent...")
        keywords = extract_search_query(title, intent)
        print(f"  Derived keywords: {keywords}")
    elif not keywords:
        keywords = title

    print(f"\n{'='*60}")
    print(f"SEO Article Writer")
    print(f"Topic   : {title}")
    print(f"Keywords: {keywords}")
    if intent:
        print(f"Intent  : {intent[:80]}{'...' if len(intent) > 80 else ''}")
    print(f"Output  : {output_dir}")
    print(f"{'='*60}")

    # Step 1: SERP Research
    research = serp_research(title, keywords, intent=intent)

    # Step 2: Refine Title
    refined_title = refine_title(title, keywords, research, intent=intent)

    # Step 3: Key Takeaways
    key_takeaways = generate_key_takeaways(refined_title, keywords, research)

    # Step 4: Outline
    outline = generate_outline(refined_title, keywords, research, key_takeaways)

    # Step 5: Write Content
    content = write_content(refined_title, keywords, outline, research, key_takeaways)

    # Step 6: Humanize
    humanized = humanize_content(content)

    # Step 7: Meta
    meta = generate_meta(refined_title, keywords, humanized)

    # Step 8: Image Search
    images = search_images(humanized, research)

    # Inject images into article
    final_article = inject_images(humanized, images)

    # Write outputs
    slug = slugify(refined_title)
    md_path, meta_path = write_outputs(slug, final_article, meta, images, output_dir, edition=edition)

    # Summary
    word_count = len(final_article.split())
    print(f"\n{'='*60}")
    print(f"DONE")
    print(f"  Title      : {refined_title}")
    print(f"  Word count : {word_count:,}")
    print(f"  Images     : {len(images)}")
    print(f"  Article    : {md_path}")
    print(f"  Meta JSON  : {meta_path}")
    print(f"{'='*60}\n")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate a full SEO article with web-sourced images."
    )
    parser.add_argument("topic", help='Article topic, e.g. "What is RAG in AI"')
    parser.add_argument(
        "--intent",
        default=None,
        help=(
            'Natural language description of what you want to achieve, e.g. '
            '"I want to explain to developers how ReAct agents work and why '
            'they are better than standard LLMs for tool use". '
            'Claude will extract the best search keywords from this description.'
        ),
    )
    parser.add_argument(
        "--keywords",
        default=None,
        help=(
            'Primary keyword(s) to target directly, e.g. "retrieval augmented generation". '
            'Use --intent instead for a richer, intent-driven search. '
            'If neither is provided, defaults to the topic.'
        ),
    )
    parser.add_argument(
        "--output-dir",
        default="./output",
        help="Directory to save the article and meta JSON (default: ./output)",
    )
    parser.add_argument(
        "--edition",
        type=int,
        default=0,
        help="Newsletter edition number shown in the author intro (e.g. --edition 31)",
    )
    args = parser.parse_args()

    if not os.getenv("ANTHROPIC_API_KEY"):
        print("ERROR: ANTHROPIC_API_KEY environment variable is not set.", file=sys.stderr)
        sys.exit(1)

    # --intent takes priority; --keywords is the legacy shorthand; topic is the fallback
    intent = args.intent or ""
    keywords = args.keywords or ("" if intent else args.topic)
    output_dir = Path(args.output_dir)

    run(title=args.topic, keywords=keywords, output_dir=output_dir, edition=args.edition, intent=intent)


if __name__ == "__main__":
    main()
